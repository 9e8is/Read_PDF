import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from io import BytesIO
from PIL import Image

def pdf_to_single_docx(pdf_path, output_docx_path, min_width=100, min_height=100):
    pdf_document = fitz.open(pdf_path)
    docx_document = Document()

    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)  # Загружаем страницу
        blocks = page.get_text("dict")["blocks"]
        
        # Множество для отслеживания уникальных xref изображений
        added_images = set()
        
        for block in blocks:
            if block["type"] == 0:  # Обработка текстовых блоков
                paragraph = docx_document.add_paragraph()
                
                for line in block["lines"]:
                    if line.get("height"):
                        paragraph.paragraph_format.line_spacing = Pt(line["height"])

                    for span in line["spans"]:
                        run = paragraph.add_run(span["text"])

                        # Применяем форматирование
                        if "bold" in span and span["bold"]:
                            run.bold = True
                        if "italic" in span and span["italic"]:
                            run.italic = True
                        if "size" in span:
                            run.font.size = Pt(span["size"])

                    # Выравнивание абзаца
                    if "align" in block:
                        if block["align"] == 1:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        elif block["align"] == 2:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        else:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            elif block["type"] == 1:  # Обработка изображений
                # Получаем все изображения на странице
                for img in page.get_images(full=True):
                    xref = img[0]
                    if xref not in added_images:  # Проверка на дублирование изображения
                        added_images.add(xref)
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]

                        # Загружаем изображение в поток и проверяем его размер
                        image_stream = BytesIO(image_bytes)
                        pil_image = Image.open(image_stream)

                        # Фильтрация изображений по ширине и высоте
                        if pil_image.width >= min_width and pil_image.height >= min_height:
                            # Дополнительные условия фильтрации (например, по формату)
                            if pil_image.format in ['JPEG', 'PNG']:  # Пример: добавлять только JPEG и PNG
                                # Конвертируем в RGB, если изображение в другом формате
                                if pil_image.mode in ("P", "RGBA"):
                                    pil_image = pil_image.convert("RGB")

                                # Добавляем новый параграф перед изображением
                                docx_document.add_paragraph()
                                image_stream.seek(0)  # Перематываем поток для вставки
                                docx_document.add_picture(image_stream, width=Inches(5))

        # Добавляем разрыв страницы после каждой страницы PDF, кроме последней
        if page_num < pdf_document.page_count - 1:
            docx_document.add_page_break()

    docx_document.save(output_docx_path)
    print(f"Файл сохранен как {output_docx_path}")

# Пример использования
pdf_path = "input.pdf"
output_docx_path = "output.docx"

pdf_to_single_docx(pdf_path, output_docx_path, min_width=200, min_height=200)
